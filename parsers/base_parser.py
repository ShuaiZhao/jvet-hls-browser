"""
Base parser class for video codec specifications.
All codec-specific parsers should inherit from this class.
"""

from abc import ABC, abstractmethod
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass
import re
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class SyntaxParameter:
    """Represents a single syntax parameter."""
    name: str
    type: str
    condition: Optional[str] = None
    semantics_ref: Optional[str] = None


@dataclass
class SyntaxStructure:
    """Represents a complete syntax structure."""
    id: str
    section: str
    name: str
    descriptor: str
    parameters: List[SyntaxParameter]


@dataclass
class SemanticInfo:
    """Represents semantic information for a parameter."""
    parameter: str
    section: str
    definition: str
    constraints: Dict
    related_parameters: List[str]


class BaseSpecParser(ABC):
    """Abstract base class for specification parsers."""

    def __init__(self, config: Dict):
        """
        Initialize parser with codec-specific configuration.

        Args:
            config: Configuration dictionary from codec_config.yaml
        """
        self.config = config
        self.doc = None
        self.syntax_structures = {}
        self.semantics = {}

    def load_document(self, docx_path: str) -> None:
        """
        Load the specification document.

        Args:
            docx_path: Path to the DOCX file
        """
        print(f"Loading document: {docx_path}")
        self.doc = Document(docx_path)
        print(f"Document loaded successfully. Total paragraphs: {len(self.doc.paragraphs)}")

    @abstractmethod
    def extract_syntax_structures(self) -> Dict[str, SyntaxStructure]:
        """
        Extract all syntax structures from Section 7.3.

        Returns:
            Dictionary mapping structure names to SyntaxStructure objects
        """
        pass

    @abstractmethod
    def extract_semantics(self) -> Dict[str, SemanticInfo]:
        """
        Extract semantic information from Section 7.4.

        Returns:
            Dictionary mapping parameter names to SemanticInfo objects
        """
        pass

    def find_section(self, section_number: str) -> List[Paragraph]:
        """
        Find all paragraphs belonging to a specific section.

        Args:
            section_number: Section number (e.g., "7.3", "7.4")

        Returns:
            List of paragraphs in that section
        """
        paragraphs = []
        in_section = False
        section_pattern = re.compile(rf'^{re.escape(section_number)}(\.\d+)*\s+')

        for para in self.doc.paragraphs:
            text = para.text.strip()

            # Check if we've entered the target section
            if section_pattern.match(text):
                in_section = True
                paragraphs.append(para)
                continue

            # Check if we've entered a different major section (e.g., from 7.3 to 7.5)
            if in_section and re.match(r'^\d+\.\d+\s+', text):
                # Check if it's still within our section
                if not text.startswith(section_number):
                    break

            if in_section:
                paragraphs.append(para)

        return paragraphs

    def find_tables_in_section(self, section_number: str) -> List[Tuple[Table, str]]:
        """
        Find all tables in a specific section with their preceding context.

        Args:
            section_number: Section number (e.g., "7.3.2")

        Returns:
            List of tuples (table, context_text)
        """
        tables_with_context = []
        current_context = ""

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                para = Paragraph(element, self.doc)
                text = para.text.strip()

                # Update context if it looks like a section header
                if re.match(rf'^{re.escape(section_number)}', text):
                    current_context = text

            elif element.tag.endswith('tbl'):
                # It's a table
                table = Table(element, self.doc)
                if section_number in current_context:
                    tables_with_context.append((table, current_context))

        return tables_with_context

    def extract_syntax_from_text(self, text: str) -> Optional[SyntaxStructure]:
        """
        Extract syntax structure from code-like text blocks.

        Args:
            text: Text containing syntax structure

        Returns:
            SyntaxStructure object or None
        """
        patterns = self.config.get('patterns', {})
        syntax_regex = patterns.get('syntax_descriptor_regex', r'(\w+)\(\s*\)\s*\{')
        param_regex = patterns.get('parameter_regex', r'(\w+)\s+(ue\(v\)|u\(\d+\)|se\(v\)|i\(\d+\)|f\(\d+\))')

        # Find structure name
        match = re.search(syntax_regex, text)
        if not match:
            return None

        structure_name = match.group(1)
        parameters = []

        # Extract parameters
        lines = text.split('\n')
        current_condition = None

        for line in lines:
            line = line.strip()

            # Check for conditions
            if 'if' in line.lower():
                # Extract condition
                cond_match = re.search(r'if\s*\((.*?)\)', line)
                if cond_match:
                    current_condition = cond_match.group(1)

            # Extract parameter
            param_match = re.search(param_regex, line)
            if param_match:
                param_name = param_match.group(1)
                param_type = param_match.group(2)

                parameters.append(SyntaxParameter(
                    name=param_name,
                    type=param_type,
                    condition=current_condition
                ))

            # Reset condition after closing brace
            if '}' in line:
                current_condition = None

        if not parameters:
            return None

        return SyntaxStructure(
            id=structure_name,
            section="",  # Will be filled by subclass
            name=structure_name,
            descriptor="",  # Will be filled by subclass
            parameters=parameters
        )

    def extract_constraints(self, text: str) -> Dict:
        """
        Extract constraints from semantic text.

        Args:
            text: Semantic description text

        Returns:
            Dictionary of constraints
        """
        constraints = {}

        # Extract range constraints
        range_pattern = r'(?:shall be|value of \w+ shall be) in the range of (\d+) to (\d+)'
        range_match = re.search(range_pattern, text)
        if range_match:
            constraints['range'] = f"{range_match.group(1)}..{range_match.group(2)}"

        # Extract value mappings (e.g., "0 indicates X, 1 indicates Y")
        value_pattern = r'(\d+)\s+(?:specifies|indicates|corresponds to)\s+([^.,\n]+)'
        value_matches = re.findall(value_pattern, text)
        if value_matches:
            constraints['values'] = {val: desc.strip() for val, desc in value_matches}

        return constraints

    def find_related_parameters(self, text: str) -> List[str]:
        """
        Find parameter names mentioned in semantic text.

        Args:
            text: Semantic description text

        Returns:
            List of related parameter names
        """
        # Common parameter naming patterns
        param_pattern = r'\b([a-z_]+(?:_[a-z0-9]+)+)\b'
        matches = re.findall(param_pattern, text)

        # Filter to likely parameter names (contain underscores, lowercase)
        related = [m for m in matches if m.count('_') >= 1]

        return list(set(related))

    def parse(self) -> Tuple[Dict, Dict]:
        """
        Parse the entire specification.

        Returns:
            Tuple of (syntax_structures, semantics)
        """
        print("Extracting syntax structures...")
        self.syntax_structures = self.extract_syntax_structures()
        print(f"Found {len(self.syntax_structures)} syntax structures")

        print("Extracting semantics...")
        self.semantics = self.extract_semantics()
        print(f"Found {len(self.semantics)} semantic definitions")

        return self.syntax_structures, self.semantics

    def to_dict(self) -> Tuple[Dict, Dict]:
        """
        Convert parsed data to dictionary format for JSON export.

        Returns:
            Tuple of (syntax_dict, semantics_dict)
        """
        syntax_dict = {}
        for name, structure in self.syntax_structures.items():
            syntax_dict[name] = {
                'id': structure.id,
                'section': structure.section,
                'name': structure.name,
                'descriptor': structure.descriptor,
                'parameters': [
                    {
                        'name': p.name,
                        'type': p.type,
                        'condition': p.condition,
                        'semantics_ref': p.semantics_ref
                    }
                    for p in structure.parameters
                ]
            }

        semantics_dict = {}
        for name, info in self.semantics.items():
            semantics_dict[name] = {
                'parameter': info.parameter,
                'section': info.section,
                'definition': info.definition,
                'constraints': info.constraints,
                'related_parameters': info.related_parameters
            }

        return syntax_dict, semantics_dict
