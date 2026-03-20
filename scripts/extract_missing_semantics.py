#!/usr/bin/env python3
"""
Extract missing semantics from H.266 VVC specification Word document
"""

import json
import re
from docx import Document
from pathlib import Path

# Paths
SPEC_PATH = "/Users/shzhao/Library/CloudStorage/GoogleDrive-dr.shuai.zhao@gmail.com/My Drive/JVET/HLS/JVET_notes/codec_specification/H266_VVC/T-REC-H.266-202309-I!!MSW-E v3.docx"
SYNTAX_JSON = "data/vvc/syntax.json"
SEMANTICS_JSON = "data/vvc/semantics.json"
OUTPUT_JSON = "data/vvc/semantics.json"

def load_existing_semantics():
    """Load existing semantics JSON"""
    with open(SEMANTICS_JSON, 'r') as f:
        return json.load(f)

def get_missing_parameters():
    """Get list of parameters in syntax but not in semantics"""
    with open(SYNTAX_JSON, 'r') as f:
        syntax_data = json.load(f)

    semantics = load_existing_semantics()

    # Extract all parameter names from syntax
    syntax_params = set()
    for structure_name, structure in syntax_data.items():
        if 'parameters' in structure:
            for param in structure['parameters']:
                if isinstance(param, dict) and param.get('condition') == 'parameter':
                    param_name = param.get('name', '').split('[')[0].strip()
                    if param_name and not param_name.startswith('//'):
                        syntax_params.add(param_name)

    missing = syntax_params - set(semantics.keys())
    return sorted(missing)

def extract_parameter_semantics(doc, param_name):
    """
    Extract semantics for a specific parameter from the Word document
    """
    paragraphs = doc.paragraphs

    # Search for the parameter
    for i, para in enumerate(paragraphs):
        text = para.text.strip()

        # Look for parameter definition (usually starts with parameter name followed by some punctuation)
        # Pattern: "param_name specifies..." or "param_name shall..." or "param_name is..."
        pattern = rf'^{re.escape(param_name)}\s+(specifies|shall|is|indicates|contains|represents|equal|has|when|provides|determines)'

        if re.match(pattern, text, re.IGNORECASE):
            # Found it! Extract this and next few paragraphs
            definition_parts = [text]

            # Continue collecting until we hit another parameter definition or empty line
            for j in range(i + 1, min(i + 15, len(paragraphs))):
                next_text = paragraphs[j].text.strip()

                # Stop if we hit another parameter definition
                if next_text and re.match(r'^[a-z_][a-z0-9_]*\s+(specifies|shall|is|indicates)', next_text, re.IGNORECASE):
                    break

                # Stop if we hit a section heading or empty lines
                if not next_text or next_text.startswith('7.') or next_text.startswith('NOTE'):
                    break

                definition_parts.append(next_text)

            definition = ' '.join(definition_parts)

            return {
                "parameter": param_name,
                "section": "",  # Could extract from context
                "definition": definition,
                "constraints": [],
                "related_parameters": []
            }

    return None

def main():
    print("=" * 80)
    print("Extracting Missing Semantics from H.266 VVC Specification")
    print("=" * 80)

    # Load document
    print(f"\nLoading specification: {SPEC_PATH}")
    doc = Document(SPEC_PATH)
    print(f"✓ Loaded {len(doc.paragraphs)} paragraphs")

    # Get missing parameters
    print("\nAnalyzing missing parameters...")
    missing_params = get_missing_parameters()
    print(f"✓ Found {len(missing_params)} missing parameters")

    # Load existing semantics
    semantics = load_existing_semantics()
    added_count = 0
    not_found_count = 0

    print("\nExtracting semantics...")
    for i, param in enumerate(missing_params, 1):
        print(f"[{i}/{len(missing_params)}] {param}...", end=' ')

        result = extract_parameter_semantics(doc, param)
        if result:
            semantics[param] = result
            added_count += 1
            print("✓")
        else:
            not_found_count += 1
            print("✗ Not found")

    print("\n" + "=" * 80)
    print(f"Results:")
    print(f"  Added: {added_count}")
    print(f"  Not found: {not_found_count}")
    print(f"  Total semantics: {len(semantics)}")
    print("=" * 80)

    # Save updated semantics
    print(f"\nSaving to {OUTPUT_JSON}...")
    with open(OUTPUT_JSON, 'w') as f:
        json.dump(semantics, f, indent=2)

    # Also update web version
    web_output = "web/data/vvc/semantics.json"
    print(f"Saving to {web_output}...")
    with open(web_output, 'w') as f:
        json.dump(semantics, f, indent=2)

    print("\n✓ Done!")

    # Print first 10 not found for debugging
    if not_found_count > 0:
        print("\nFirst 10 parameters not found:")
        count = 0
        for param in missing_params:
            if param not in semantics:
                print(f"  - {param}")
                count += 1
                if count >= 10:
                    break

if __name__ == "__main__":
    main()
